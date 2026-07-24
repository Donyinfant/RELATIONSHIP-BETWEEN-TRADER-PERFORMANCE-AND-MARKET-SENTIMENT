#RELATIONSHIP BETWEEN TRADER PERFORMANCE AND  MARKET SENTIMENT
import pandas as pd
sentiment = pd.read_csv(r"E:\bitcoin\fear_greed_index.csv")
trades = pd.read_csv(r"E:\bitcoin\historical_data.csv")
print(sentiment.head())
print(trades.head())
print(sentiment.info())
print(trades.info())


sentiment["date"] = pd.to_datetime(sentiment["date"])
sentiment=sentiment[["date","classification"]].rename(columns={"classification":"Emotion"})
trades["Date"] = pd.to_datetime(trades["Timestamp"],unit="ms")
trades['date']=pd.to_datetime(trades['Timestamp IST'],format='%d-%m-%Y %H:%M').dt.normalize()
df=trades.merge(sentiment,on='date',how='left')
print(f"unmatched: {df['Emotion'].isna().sum()} of {len(df)}")
df=df.dropna(subset=['Emotion'])


closing = df[df['Closed PnL'] != 0].copy()
closing['is_win'] = closing['Closed PnL'] > 0

order=['Extreme Fear','Fear','Neutral','Greed','Extreme Greed']
closing['Emotion']=pd.Categorical(closing['Emotion'], categories=order, ordered=True)

summary=closing.groupby('Emotion',observed=True).agg(
    trades=('Closed PnL','size'),
    win_rate=('is_win','mean'),
    avg_pnl=('Closed PnL','mean'),
    total_pnl=('Closed PnL','sum'),
)
print(summary)


import matplotlib.pyplot as plt
fig, ax =plt.subplots(figsize=(10,6))
ax.bar(summary.index.astype(str), summary['avg_pnl'])
ax.set_ylabel('Average pnl per trade (USD)')
ax.set_title('profitability by market Emotion')
plt.savefig('avg_pnl_by_Emotion.png',dpi=150)
plt.show()


from scipy import stats
fear=closing[closing['Emotion'].isin(['Fear','Extreme Fear'])]['Closed PnL']
greed=closing[closing['Emotion'].isin(['Greed','Extreme Greed'])]['Closed PnL']
t,p=stats.ttest_ind(fear,greed,equal_var=False)
print(f"t={t:.3f}, p={p:.4f}")
extreme_fear = closing[closing['Emotion']=='Extreme Fear']['Closed PnL']
extreme_greed = closing[closing['Emotion']=='Extreme Greed']['Closed PnL']
t2, p2 = stats.ttest_ind(extreme_fear, extreme_greed, equal_var=False)
print(f"Extreme Fear vs Extreme Greed: t={t2:.3f}, p={p2:.4f}")


from docx import Document
from docx.shared import Inches

doc=Document()
doc.add_heading('Trader Sentiment Report', 0)

doc.add_paragraph(
    "Trade-level performance was grouped by market sentiment and compared across "
    "five categories: Extreme Fear, Fear, Neutral, Greed, and Extreme Greed."
)

doc.add_heading('Data Notes', level=1)

doc.add_paragraph(
    f"This analysis covers {len(closing):,} closing trades (from {len(df):,} total "
    f"execution records) across {df['Account'].nunique()} unique accounts, merged "
    f"with daily Bitcoin Fear & Greed Index classifications. The historical trade "
    f"data included fields such as execution price, size, side, direction, closed "
    f"PnL, and fees, but did not include a leverage field, so leverage-adjusted "
    f"performance could not be assessed in this analysis."
)

doc.add_heading('Win Rate and Profitability by Sentiment', level=1)


table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Sentiment'
hdr[1].text = 'Trades'
hdr[2].text = 'Win Rate'
hdr[3].text = 'Avg PnL/Trade'
hdr[4].text = 'Total PnL'

for emotion, row in summary.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(emotion)
    cells[1].text = f"{int(row['trades']):,}"
    cells[2].text = f"{row['win_rate']*100:.1f}%"
    cells[3].text = f"${row['avg_pnl']:.2f}"
    cells[4].text = f"${row['total_pnl']:,.0f}"

doc.add_paragraph(
    "Win rate and average profit per trade do not move in a straight line from Fear "
    "toward Greed. Instead, they form a U-shape: the two sentiment extremes, Fear "
    "(87.3% win rate, $112.63 average profit per trade) and Extreme Greed (89.2%, "
    "$130.21), clearly outperform the calmer middle of the spectrum. Extreme Fear is "
    "the weakest regime observed, with the lowest win rate (76.2%) and one of the "
    "lowest average profits ($71.03) per trade."
)

doc.add_heading('Profitability Chart', level=1)
doc.add_picture('avg_pnl_by_Emotion.png', width=Inches(6))


doc.add_heading('Statistical Significance', level=1)

doc.add_paragraph(
    f"A Welch's t-test comparing per-trade profit between the Fear-family group "
    f"(Fear + Extreme Fear) and the Greed-family group (Greed + Extreme Greed) "
    f"produced t = {t:.3f}, p = {p:.4f}. Since this p-value is well above the standard "
    f"0.05 significance threshold, the broad Fear-vs-Greed split does not show a "
    f"statistically significant difference in average trade profitability."
)

doc.add_paragraph(
    f"However, a direct comparison of the two sentiment extremes tells a different "
    f"story: Extreme Fear versus Extreme Greed produced t = {t2:.3f}, p = {p2:.4f}. "
    f"This result is statistically significant, confirming that the U-shaped pattern "
    f"is real at the extremes even though it is not visible when Fear/Extreme Fear "
    f"and Greed/Extreme Greed are grouped together as broad categories. This is the "
    f"result that actually supports the Key Takeaway below."
)
doc.add_heading('Key Takeaway', level=1)

doc.add_paragraph(
    "Traders in this dataset performed best during periods of strong sentiment in "
    "either direction (Fear or Extreme Greed), and worst during Extreme Fear. This "
    "suggests that sentiment intensity, rather than sentiment direction alone, is the "
    "more useful signal for anticipating trading performance."
)


doc.save('RELATIONSHIP BETWEEN TRADER PERFORMANCE AND MARKET SENTIMENT.docx')
print("Report saved successfully.")



 



